"""
Export Manager Module - Handles exporting text to various formats.
Supports TXT, Markdown, and DOCX formats.
"""

import os
from pathlib import Path
from typing import Optional
from dataclasses import dataclass
from datetime import datetime

# Try to import python-docx for DOCX export
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ExportResult:
    """Result of export operation."""
    success: bool
    file_path: str
    message: str
    format_type: str


class ExportManager:
    """Manages exporting text content to various formats."""

    def __init__(self):
        self.docx_available = DOCX_AVAILABLE

    def export(self, text: str, file_path: str, format_type: str,
               title: Optional[str] = None,
               metadata: Optional[dict] = None) -> ExportResult:
        """
        Export text to specified format.

        Args:
            text: Text content to export
            file_path: Output file path
            format_type: One of 'txt', 'md', 'docx'
            title: Optional document title
            metadata: Optional metadata dict

        Returns:
            ExportResult with success status
        """
        format_type = format_type.lower()

        if format_type == 'txt':
            return self._export_txt(text, file_path, title, metadata)
        elif format_type in ('md', 'markdown'):
            return self._export_markdown(text, file_path, title, metadata)
        elif format_type == 'docx':
            return self._export_docx(text, file_path, title, metadata)
        else:
            return ExportResult(
                success=False,
                file_path=file_path,
                message=f"Unsupported format: {format_type}",
                format_type=format_type
            )

    def _export_txt(self, text: str, file_path: str,
                    title: Optional[str], metadata: Optional[dict]) -> ExportResult:
        """Export as plain text."""
        try:
            content = ""

            if title:
                content += f"{title}\n{'=' * len(title)}\n\n"

            if metadata:
                content += f"Generated: {metadata.get('timestamp', datetime.now().isoformat())}\n"
                if 'mode' in metadata:
                    content += f"Mode: {metadata['mode']}\n"
                if 'creativity' in metadata:
                    content += f"Creativity: {metadata['creativity']}%\n"
                content += "\n---\n\n"

            content += text

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return ExportResult(
                success=True,
                file_path=file_path,
                message=f"Exported to {file_path}",
                format_type='txt'
            )

        except Exception as e:
            return ExportResult(
                success=False,
                file_path=file_path,
                message=f"Export failed: {str(e)}",
                format_type='txt'
            )

    def _export_markdown(self, text: str, file_path: str,
                         title: Optional[str], metadata: Optional[dict]) -> ExportResult:
        """Export as Markdown."""
        try:
            content = ""

            if title:
                content += f"# {title}\n\n"

            if metadata:
                content += "> **Enhanced by LocalWrite**  \n"
                content += f"> Date: {metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M'))}  \n"
                if 'mode' in metadata:
                    content += f"> Mode: {metadata['mode']}  \n"
                if 'creativity' in metadata:
                    content += f"> Creativity: {metadata['creativity']}%  \n"
                if 'ai_score' in metadata:
                    content += f"> AI Score: {metadata['ai_score']}%  \n"
                content += "\n---\n\n"

            # Convert text to markdown paragraphs
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    content += para.strip() + "\n\n"

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return ExportResult(
                success=True,
                file_path=file_path,
                message=f"Exported to {file_path}",
                format_type='md'
            )

        except Exception as e:
            return ExportResult(
                success=False,
                file_path=file_path,
                message=f"Export failed: {str(e)}",
                format_type='md'
            )

    def _export_docx(self, text: str, file_path: str,
                     title: Optional[str], metadata: Optional[dict]) -> ExportResult:
        """Export as DOCX (Word document)."""
        if not DOCX_AVAILABLE:
            return ExportResult(
                success=False,
                file_path=file_path,
                message="DOCX export requires python-docx. Install with: pip install python-docx",
                format_type='docx'
            )

        try:
            doc = Document()

            # Add title
            if title:
                heading = doc.add_heading(title, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            if metadata:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = meta_para.add_run(f"Enhanced by LocalWrite")
                run.italic = True
                run.font.size = Pt(10)

                if 'timestamp' in metadata:
                    meta_para.add_run(f" | {metadata['timestamp']}")
                if 'mode' in metadata:
                    meta_para.add_run(f" | Mode: {metadata['mode']}")
                if 'creativity' in metadata:
                    meta_para.add_run(f" | Creativity: {metadata['creativity']}%")

                doc.add_paragraph()  # Spacer

            # Add main content
            paragraphs = text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.paragraph_format.space_after = Pt(12)

            # Save document
            doc.save(file_path)

            return ExportResult(
                success=True,
                file_path=file_path,
                message=f"Exported to {file_path}",
                format_type='docx'
            )

        except Exception as e:
            return ExportResult(
                success=False,
                file_path=file_path,
                message=f"Export failed: {str(e)}",
                format_type='docx'
            )

    def export_comparison(self, original: str, humanized: str, file_path: str,
                          format_type: str, metadata: Optional[dict] = None) -> ExportResult:
        """Export original and humanized text side by side."""
        if format_type == 'txt':
            content = "ORIGINAL TEXT\n" + "=" * 50 + "\n\n"
            content += original + "\n\n"
            content += "HUMANIZED TEXT\n" + "=" * 50 + "\n\n"
            content += humanized
            return self._export_txt(content, file_path, "Text Comparison", metadata)

        elif format_type in ('md', 'markdown'):
            content = "## Original Text\n\n"
            content += original + "\n\n"
            content += "---\n\n"
            content += "## Humanized Text\n\n"
            content += humanized
            return self._export_markdown(content, file_path, "Text Comparison", metadata)

        elif format_type == 'docx':
            if not DOCX_AVAILABLE:
                return ExportResult(
                    success=False,
                    file_path=file_path,
                    message="DOCX export requires python-docx",
                    format_type='docx'
                )

            try:
                doc = Document()
                doc.add_heading("Text Comparison", level=0)

                if metadata:
                    meta = doc.add_paragraph()
                    meta.add_run(f"Generated: {metadata.get('timestamp', '')}").italic = True

                doc.add_heading("Original Text", level=1)
                for para in original.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())

                doc.add_page_break()

                doc.add_heading("Humanized Text", level=1)
                for para in humanized.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())

                doc.save(file_path)

                return ExportResult(
                    success=True,
                    file_path=file_path,
                    message=f"Exported comparison to {file_path}",
                    format_type='docx'
                )

            except Exception as e:
                return ExportResult(
                    success=False,
                    file_path=file_path,
                    message=f"Export failed: {str(e)}",
                    format_type='docx'
                )

        return ExportResult(
            success=False,
            file_path=file_path,
            message=f"Unsupported format: {format_type}",
            format_type=format_type
        )

    def get_available_formats(self) -> list:
        """Get list of available export formats."""
        formats = [
            {'id': 'txt', 'name': 'Plain Text', 'ext': '.txt', 'available': True},
            {'id': 'md', 'name': 'Markdown', 'ext': '.md', 'available': True},
            {'id': 'docx', 'name': 'Word Document', 'ext': '.docx', 'available': DOCX_AVAILABLE}
        ]
        return formats

    def is_format_available(self, format_type: str) -> bool:
        """Check if format is available."""
        if format_type in ('txt', 'md', 'markdown'):
            return True
        elif format_type == 'docx':
            return DOCX_AVAILABLE
        return False


# Singleton instance
_export_manager = ExportManager()


def export_text(text: str, file_path: str, format_type: str,
                title: Optional[str] = None,
                metadata: Optional[dict] = None) -> ExportResult:
    """Convenience function to export text."""
    return _export_manager.export(text, file_path, format_type, title, metadata)


def get_export_manager() -> ExportManager:
    """Get the global export manager instance."""
    return _export_manager


def is_docx_available() -> bool:
    """Check if DOCX export is available."""
    return DOCX_AVAILABLE
